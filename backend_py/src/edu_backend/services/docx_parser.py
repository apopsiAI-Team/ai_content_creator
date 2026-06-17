"""
DOCX educational-design parser — Python port of web/src/utils/docxParser.ts

Extracts structured data (title, total hours, modules with ESCO skills) from a
.docx educational design document. Mirrors the regex-based logic of the
frontend mammoth-based parser so that API consumers get identical results.

Keep this file in sync with web/src/utils/docxParser.ts when the document
format changes.
"""
from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from typing import Literal

from docx import Document


SkillType = Literal["essential", "optional"]


@dataclass
class ESCOSkill:
    code: str
    name: str
    type: SkillType


@dataclass
class Module:
    number: int
    title: str
    hours: int
    content: str = ""
    activities: str = ""
    skills: list[ESCOSkill] = field(default_factory=list)


@dataclass
class ParsedDocument:
    documentTitle: str
    totalHours: int
    modules: list[Module]


def parse_docx_bytes(data: bytes) -> ParsedDocument:
    """Parse a .docx file (bytes) into a ParsedDocument."""
    doc = Document(io.BytesIO(data))
    text = "\n".join(p.text for p in doc.paragraphs)
    return parse_educational_design(text)


def parse_educational_design(text: str) -> ParsedDocument:
    lines = [ln for ln in text.split("\n") if ln.strip()]

    title = ""
    occupation_match = re.search(
        r"(?:για\s+)?(στελέχη[^,\n.]+|Ειδικός[^,\n.]+)", text, re.IGNORECASE
    )
    if occupation_match:
        occupation = occupation_match.group(1).strip()
        title = f"Πρόγραμμα επαγγελματικής κατάρτισης για {occupation}"
    else:
        for line in lines:
            if "Ειδικός" in line or "Πρόγραμμα" in line:
                short_match = re.search(
                    r"(Ειδικός[^,\n]+|Πρόγραμμα[^για]*για[^,\n]+)",
                    line,
                    re.IGNORECASE,
                )
                title = short_match.group(1).strip() if short_match else line.strip()[:80]
                break

    if not title and lines:
        title = lines[0].strip()[:80]

    total_hours = 0
    hours_match = re.search(r"Συνολικές\s*[Ωώ]ρες[:\s]*(\d+)", text, re.IGNORECASE)
    if hours_match:
        total_hours = int(hours_match.group(1))

    modules = _extract_modules(text)

    if not total_hours:
        total_hours = sum(m.hours for m in modules)

    return ParsedDocument(documentTitle=title, totalHours=total_hours, modules=modules)


def _extract_modules(text: str) -> list[Module]:
    thematic_start = text.find("Θεματικές Ενότητες")
    working_text = text[thematic_start:] if thematic_start >= 0 else text

    module_pattern = re.compile(
        r"(\d+)\.\s+([^\n]+)\n[^\d]*?(\d+)\s*ώρες", re.DOTALL
    )
    matches = list(module_pattern.finditer(working_text))

    if not matches:
        return _extract_modules_alternative(text)

    modules: list[Module] = []
    for i, match in enumerate(matches):
        number = int(match.group(1))
        m_title = match.group(2).strip()
        hours = int(match.group(3))

        start_idx = match.end()
        end_idx = matches[i + 1].start() if i + 1 < len(matches) else len(working_text)
        section_text = working_text[start_idx:end_idx]

        content_match = re.search(
            r"Περιεχόμενο:\s*([^🎮Δ]+?)(?=Δραστηριότητες|ESCO|$)",
            section_text,
            re.DOTALL,
        )
        content = content_match.group(1).strip() if content_match else ""

        activities_match = re.search(
            r"Δραστηριότητες:\s*([^🎯E]+?)(?=ESCO|$)",
            section_text,
            re.DOTALL,
        )
        activities = activities_match.group(1).strip() if activities_match else ""

        skills_match = re.search(
            r"ESCO\s*Skills[^:]*:\s*([\s\S]*?)(?=\d+\.|$)",
            section_text,
        )
        skills_text = skills_match.group(1) if skills_match else ""
        skills = _parse_skills(skills_text)

        modules.append(
            Module(
                number=number,
                title=m_title,
                hours=hours,
                content=content,
                activities=activities,
                skills=skills,
            )
        )

    modules.sort(key=lambda m: m.number)
    return modules


def _extract_modules_alternative(text: str) -> list[Module]:
    modules: list[Module] = []
    lines = text.split("\n")
    current: Module | None = None
    content_buffer = ""

    module_header = re.compile(r"^(\d+)\.\s+(.+?)(?:\s*[-–]\s*(\d+)\s*ώρες)?$")
    next_line_hours = re.compile(r"^\s*(\d+)\s*ώρες", re.IGNORECASE)
    inline_hours = re.compile(r"(\d+)\s*ώρες", re.IGNORECASE)

    for i, raw_line in enumerate(lines):
        line = raw_line.strip()
        next_line = lines[i + 1].strip() if i + 1 < len(lines) else ""

        m = module_header.match(line)
        if m:
            has_hours_in_line = m.group(3) is not None
            next_match = next_line_hours.match(next_line)
            has_hours_in_next_line = next_match is not None
            lower = line.lower()
            has_module_keyword = (
                "ενότητα" in lower or "θεματική" in lower or "κεφάλαιο" in lower
            )

            if not has_hours_in_line and not has_hours_in_next_line and not has_module_keyword:
                if current is not None:
                    content_buffer += line + "\n"
                continue

            if current is not None and current.number and current.title:
                current.content = content_buffer.strip()
                modules.append(current)

            if has_hours_in_line:
                hours = int(m.group(3))
            elif next_match:
                hours = int(next_match.group(1))
            else:
                hours = 0

            current = Module(
                number=int(m.group(1)),
                title=m.group(2).strip(),
                hours=hours,
            )
            content_buffer = ""
            continue

        if current is not None:
            ih = inline_hours.search(line)
            if ih and not current.hours:
                current.hours = int(ih.group(1))

            if current is not None and ("ESCO" in line or "[E]" in line or "[O]" in line):
                current.skills.extend(_parse_skills(line))

            content_buffer += line + "\n"

    if current is not None and current.number and current.title:
        current.content = content_buffer.strip()
        modules.append(current)

    return modules


def _parse_skills(text: str) -> list[ESCOSkill]:
    skills: list[ESCOSkill] = []

    skill_pattern = re.compile(r"([^\[\]\n,]+)\s*\[([EO])\]")
    for match in skill_pattern.finditer(text):
        name = match.group(1).strip()
        skill_type: SkillType = "essential" if match.group(2) == "E" else "optional"
        if name and len(name) > 2:
            skills.append(ESCOSkill(code=_generate_skill_code(name), name=name, type=skill_type))

    if not skills:
        for part in re.split(r"[,\n]", text):
            name = part.strip()
            if name and len(name) > 3 and not re.match(r"^\d", name) and "ESCO" not in name:
                skills.append(
                    ESCOSkill(code=_generate_skill_code(name), name=name, type="essential")
                )

    return skills


def _generate_skill_code(name: str) -> str:
    cleaned = re.sub(r"[^\w\sα-ωά-ώ]", "", name.lower(), flags=re.UNICODE)
    parts = re.split(r"\s+", cleaned)
    code = "-".join(p[:3] for p in parts if p)
    return code[:20]


def to_dict(parsed: ParsedDocument) -> dict:
    """Serialize ParsedDocument to a JSON-friendly dict matching the frontend shape."""
    return {
        "documentTitle": parsed.documentTitle,
        "totalHours": parsed.totalHours,
        "modules": [
            {
                "number": m.number,
                "title": m.title,
                "hours": m.hours,
                "content": m.content,
                "activities": m.activities,
                "skills": [
                    {"code": s.code, "name": s.name, "type": s.type} for s in m.skills
                ],
            }
            for m in parsed.modules
        ],
    }
